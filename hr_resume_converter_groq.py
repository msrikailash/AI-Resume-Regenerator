import os
import tempfile
import re
import pythoncom
import pdfplumber
from datetime import datetime
from flask import Flask, request, render_template, send_file, flash, redirect, url_for, after_this_request
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
from groq import Groq

# ================= CONFIGURATION =================
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = "llama-3.3-70b-versatile"

app = Flask(__name__)
app.secret_key = "hr-resume-converter-2026"
TEMP_DIR = tempfile.gettempdir()

# ================= HELPERS =================
def safe_convert_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF safely on Windows"""
    try:
        pythoncom.CoInitialize()
        convert(docx_path, pdf_path)
        return pdf_path
    except Exception as e:
        print(f"PDF Error: {e}")
        return docx_path
    finally:
        pythoncom.CoUninitialize()

def extract_text(path):
    """Extract text from PDF or DOCX"""
    text = ""
    try:
        if path.lower().endswith(".pdf"):
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
        else:
            doc = Document(path)
            text = "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"Extraction Error: {e}")
    return text.strip()

def get_ai_data(resume_text):
    """Get structured data from Groq AI"""
    prompt = f"""Extract professional details from this resume.
STRICT RULES:
1. ONLY use details from the text.
2. If missing, leave empty.
3. Separate Technical and Soft skills.
4. Professional Title should be JUST the job title (e.g. "Java Developer").

FORMAT:
CANDIDATE INFORMATION:
- Full Name:
- Professional Title:
- Email:
- Phone:
- Location:

PROFILE SUMMARY:
(Short summary)

PROFESSIONAL EXPERIENCE:
(Format: Company | Role | Duration | Responsibilities)

PROJECT EXPERIENCE:
(Projects and tech used)

TECHNICAL SKILLS:
(Tools, languages, etc.)

SOFT SKILLS:
(Communication, etc.)

TEXT:
{resume_text[:4000]}
"""
    try:
        client = Groq(api_key=GROQ_API_KEY)
        resp = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        return parse_ai_response(resp.choices[0].message.content)
    except Exception as e:
        print(f"AI Error: {e}")
        return None

def parse_ai_response(text):
    """Parse AI output into dictionary"""
    data = {
        "Full Name": "", "Professional Title": "", "Email": "", "Phone": "", "Location": "",
        "Profile Summary": "", "Professional Experience": "", "Project Experience": "",
        "Technical Skills": "", "Soft Skills": ""
    }
    
    current_key = None
    for line in text.split("\n"):
        line = line.strip()
        if not line: continue
        
        lower_line = line.lower()
        if "full name:" in lower_line: data["Full Name"] = line.split(":", 1)[1].strip(); continue
        if "professional title:" in lower_line: data["Professional Title"] = line.split(":", 1)[1].strip(); continue
        if "email:" in lower_line: data["Email"] = line.split(":", 1)[1].strip(); continue
        if "phone:" in lower_line: data["Phone"] = line.split(":", 1)[1].strip(); continue
        if "location:" in lower_line: data["Location"] = line.split(":", 1)[1].strip(); continue
        
        if "profile summary:" in lower_line: current_key = "Profile Summary"; continue
        if "professional experience:" in lower_line: current_key = "Professional Experience"; continue
        if "project experience:" in lower_line: current_key = "Project Experience"; continue
        if "technical skills:" in lower_line: current_key = "Technical Skills"; continue
        if "soft skills:" in lower_line: current_key = "Soft Skills"; continue
        
        if current_key:
            data[current_key] += line + "\n"
            
    return {k: v.strip() for k, v in data.items()}

def create_resume(data, path):
    """Generate Word document from data"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(10)
    
    # Page Setup & Border
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.8)
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for b in ('top', 'left', 'bottom', 'right'):
            el = OxmlElement(f'w:{b}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '8')
            el.set(qn('w:space'), '24')
            el.set(qn('w:color'), '2980B9')
            pgBorders.append(el)
        sectPr.append(pgBorders)

    # Logo
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    static_dir = os.path.join(os.path.dirname(__file__), 'static')
    logo_found = False
    for ext in ['png', 'jpg', 'jpeg', 'bmp']:
        lp = os.path.join(static_dir, f'krify_logo.{ext}')
        if os.path.exists(lp):
            logo_para.add_run().add_picture(lp, width=Inches(1.0))
            logo_found = True; break
    if not logo_found:
        r = logo_para.add_run("KRIFY")
        r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = RGBColor(0, 55, 101)
    logo_para.space_after = Pt(12)

    # Name and Title
    name = data["Full Name"].upper()
    title = re.sub(r'^(PROFESSIONAL TITLE/DESIGNATION:|TITLE:|DESIGNATION:)\s*', '', data["Professional Title"], flags=re.I).strip().upper()
    display = f"{name} – {title}" if title else name
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(display)
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = RGBColor(0, 55, 101)
    p.space_after = Pt(10)

    # Contact
    parts = [v for k, v in data.items() if k in ["Email", "Phone", "Location"] and v]
    if parts:
        p = doc.add_paragraph(" | ".join(parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = RGBColor(127, 140, 141)
        p.space_after = Pt(20)

    # Sections
    sections = [
        ("PROFILE SUMMARY", "Profile Summary"),
        ("PROFESSIONAL EXPERIENCE", "Professional Experience"),
        ("PROJECT EXPERIENCE", "Project Experience"),
        ("TECHNICAL SKILLS", "Technical Skills"),
        ("SOFT SKILLS", "Soft Skills")
    ]
    
    for title, key in sections:
        if data[key]:
            tp = doc.add_paragraph()
            tr = tp.add_run(title)
            tr.font.size = Pt(13); tr.font.bold = True; tr.font.color.rgb = RGBColor(0, 55, 101)
            tp.space_before = Pt(14); tp.space_after = Pt(8)
            
            for line in data[key].split("\n"):
                if not line.strip(): continue
                cp = doc.add_paragraph(line.strip())
                cp.runs[0].font.color.rgb = RGBColor(52, 73, 94)
                if line.strip().startswith(("-", "•")):
                    cp.style = 'List Bullet'
                    cp.paragraph_format.left_indent = Inches(0.25)

    # Footer
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph("Private and Confidential Document – Intended for authorized organizations only.")
    p.runs[0].font.size = Pt(9); p.runs[0].font.bold = True; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph("info@krify.com")
    p.runs[0].font.size = Pt(9); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.save(path)
    return safe_convert_to_pdf(path, path.replace(".docx", ".pdf"))

@app.route("/")
def index(): return render_template("hr_converter.html")

@app.route("/convert", methods=["POST"])
def convert_route():
    file = request.files.get("candidate_resume")
    if not file: return redirect("/")
    
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    in_path = os.path.join(TEMP_DIR, f"in_{ts}_{file.filename}")
    file.save(in_path)
    
    cleanup_list = [in_path]
    try:
        txt = extract_text(in_path)
        data = get_ai_data(txt)
        if not data: return redirect("/")
        
        safe_name = re.sub(r'[^a-zA-Z0-9]', '_', data["Full Name"])[:30] or "Candidate"
        out_path = os.path.join(TEMP_DIR, f"Resume_{safe_name}_{ts}.docx")
        final_path = create_resume(data, out_path)
        
        cleanup_list.extend([out_path, final_path])
        
        @after_this_request
        def cleanup(resp):
            for f in cleanup_list:
                if os.path.exists(f): os.remove(f)
            return resp
            
        return send_file(final_path, as_attachment=True, download_name=os.path.basename(final_path))
    except Exception as e:
        print(f"Route Error: {e}")
        return redirect("/")

if __name__ == "__main__":
    app.run(debug=True, port=5003)
